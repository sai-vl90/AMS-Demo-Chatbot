import os
import json
import re
import unicodedata
import ftfy
import pandas as pd
import numpy as np
import logging
from typing import List, Dict, Tuple, Set, Optional
from dataclasses import dataclass, asdict

# For DOCX parsing
from docx import Document
from docx.document import Document as _Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

# This lets us inspect runs for manual page breaks (w:br w:type="page")
from docx.enum.text import WD_BREAK

# For PDF parsing
from PyPDF2 import PdfReader

from langchain.text_splitter import RecursiveCharacterTextSplitter
from sklearn.feature_extraction.text import TfidfVectorizer

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@dataclass
class SectionMetadata:
    """Metadata for document sections"""
    title: str
    level: int
    path: str
    sap_module: Optional[str] = None
    transaction_codes: List[str] = None
    wricef_items: List[str] = None
    page_number: Optional[int] = None

@dataclass
class TableMetadata:
    """Metadata for tables in document"""
    table_type: str  # 'configuration', 'master_data', 'process', etc.
    headers: List[str]
    context: str
    row_count: int
    col_count: int
    location: Dict[str, int]  # e.g. {"table_index": 0}
    referenced_by: List[str]  # sections referencing this table

class DocumentPreprocessor:
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        # Initialize TF-IDF Vectorizer for keyword extraction
        self.vectorizer = TfidfVectorizer(
            max_features=20, 
            stop_words='english',
            ngram_range=(1, 2)  # Include bigrams for better SAP terminology
        )
        
        # SAP-specific patterns
        self.sap_patterns = {
            'transaction_codes': r't-code:?\s*([A-Z0-9]{2,8})',
            'module_codes': r'\b(FI|CO|MM|SD|PP|QM|PM|HR|PS)-\d+\b',
            'wricef_items': r'\b(RE|WF|IN|CO|EN|FR)-\d+\b',
            'config_tables': r'\b[A-Z]{4}\b'
        }
        
        # Will be set when processing DOCX
        self.doc = None

    ################################################################
    #                         CLEAN TEXT
    ################################################################
    def clean_text(self, text: str) -> str:
        """Enhanced text cleaning with SAP-specific handling."""
        if not text:
            return ""
            
        # Fix text encoding issues
        text = ftfy.fix_text(text)
        text = unicodedata.normalize("NFKC", text)
        
        # Remove bracketed/curly content (e.g., tracked changes/comments)
        text = re.sub(r'\[.*?\]|\{.*?\}', '', text)
        
        # Remove special Unicode chars (non-breaking spaces, zero-width, etc.)
        text = re.sub(r'[\u00A0\u200B\u200C\u200D\u2060\uFEFF]', ' ', text)
        
        # Normalize SAP-specific content
        text = self._normalize_sap_content(text)
        
        # Collapse extra whitespace
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    def _normalize_sap_content(self, text: str) -> str:
        """Normalize SAP-specific content (transaction codes, modules, etc.)."""
        # Normalize transaction codes
        text = re.sub(
            r't-code\s*:\s*(\w+)', 
            lambda m: f"T-Code: {m.group(1).upper()}", 
            text, 
            flags=re.IGNORECASE
        )
        
        # Normalize module codes (e.g., FI-123 → FI-123)
        text = re.sub(
            r'\b(FI|CO|MM|SD|PP|QM|PM|HR|PS)-\d+', 
            lambda m: m.group().upper(), 
            text
        )
        
        # Normalize WRICEF references (e.g., RE-123 → RE-123)
        text = re.sub(
            r'\b(RE|WF|IN|CO|EN|FR)-\d+', 
            lambda m: m.group().upper(), 
            text
        )
        
        return text

    def extract_sap_objects(self, text: str) -> Dict[str, List[str]]:
        """Extract T-Codes, WRICEF items, etc. from text."""
        return {
            'transaction_codes': re.findall(self.sap_patterns['transaction_codes'], text, re.I),
            'module_codes': re.findall(self.sap_patterns['module_codes'], text, re.I),
            'wricef_items': re.findall(self.sap_patterns['wricef_items'], text, re.I),
            'config_tables': re.findall(self.sap_patterns['config_tables'], text)
        }

    ################################################################
    #                       DOCX PARSING
    ################################################################
    def extract_sections_docx(self, filepath: str) -> List[Dict]:
        """
        Extract hierarchical sections from a DOCX file, 
        *including* manual page-break detection for page_number.
        """
        self.doc = Document(filepath)
        sections = []
        current_path = []
        current_section = None
        
        # We'll track the 'current_page' based on manual breaks
        current_page = 1

        for para in self.doc.paragraphs:
            # Check each run for a manual page break
            for run in para.runs:
                # The XML for a page break is <w:br w:type="page"/>
                if "<w:br w:type=\"page\"/>" in run.element.xml:
                    current_page += 1

            # If it's a heading (e.g. "Heading 1", "Heading 2", etc.)
            if para.style and para.style.name.startswith('Heading'):
                level_str = para.style.name.replace('Heading ', '')
                try:
                    level = int(level_str)
                except ValueError:
                    level = 1  # fallback if it's some custom style name

                title = self.clean_text(para.text)
                
                # Update section path (e.g., parent → child headings)
                while len(current_path) >= level:
                    current_path.pop()
                current_path.append(title)
                
                # Extract SAP-specific objects
                sap_objects = self.extract_sap_objects(para.text)
                
                # Create section metadata
                metadata = SectionMetadata(
                    title=title,
                    level=level,
                    path=' > '.join(current_path),
                    sap_module=self._extract_sap_module(para.text),
                    transaction_codes=sap_objects['transaction_codes'],
                    wricef_items=sap_objects['wricef_items'],
                    page_number=current_page
                )
                
                # Start a new section dictionary
                current_section = {
                    'title': title,
                    'level': level,
                    'content': '',
                    'metadata': asdict(metadata)
                }
                sections.append(current_section)
            
            elif current_section is not None and para.text.strip():
                # Append paragraph text to current section
                current_section['content'] += self.clean_text(para.text) + '\n'

        return sections

    def extract_tables_docx(self, filepath: str) -> Tuple[List[Dict], List[str]]:
        """Extract tables (with minimal context) from a DOCX file."""
        if not self.doc:
            self.doc = Document(filepath)
            
        tables = []
        table_contexts = []
        
        for idx, table in enumerate(self.doc.tables):
            try:
                table_type = self._classify_table(table)
                table_content = self._extract_table_content(table)
                if not table_content.strip():
                    continue
                
                context = self._extract_table_context(table, idx)
                metadata = TableMetadata(
                    table_type=table_type,
                    headers=self._extract_header_row(table),
                    context=context,
                    row_count=len(table.rows),
                    col_count=len(table.columns),
                    location={'table_index': idx},
                    referenced_by=[]
                )
                
                tables.append({
                    'content': table_content,
                    'metadata': asdict(metadata)
                })
                table_contexts.append(context)
                
            except Exception as e:
                self.logger.warning(f"Error processing table {idx} in DOCX: {str(e)}")
                tables.append(None)
                table_contexts.append("Error processing table")
                
        return tables, table_contexts

    ################################################################
    #                       PDF PARSING
    ################################################################
    def extract_sections_pdf(self, filepath: str) -> List[Dict]:
        """
        Very basic PDF "section" parsing: treat each page as a single "section".
        If you need more structured headings, consider a more advanced library 
        like 'unstructured' or 'pdfplumber' to detect headings/tables, etc.
        """
        sections = []
        try:
            pdf_reader = PdfReader(filepath)
            for i, page in enumerate(pdf_reader.pages):
                raw_text = page.extract_text() or ""
                clean_pg_text = self.clean_text(raw_text)
                
                # Extract SAP objects
                sap_objects = self.extract_sap_objects(clean_pg_text)

                metadata = SectionMetadata(
                    title=f"Page {i+1}",
                    level=1,
                    path=f"Page {i+1}",
                    sap_module=None,  # or attempt to parse from text
                    transaction_codes=sap_objects['transaction_codes'],
                    wricef_items=sap_objects['wricef_items'],
                    page_number=i + 1
                )

                section = {
                    'title': f"Page {i+1}",
                    'level': 1,
                    'content': clean_pg_text,
                    'metadata': asdict(metadata)
                }
                sections.append(section)
        except Exception as e:
            self.logger.error(f"Error extracting PDF sections: {str(e)}")
        return sections

    def extract_tables_pdf(self, filepath: str) -> Tuple[List[Dict], List[str]]:
        """
        Stub for PDF table extraction. 
        If you need real PDF table logic, use 'unstructured' or 'pdfplumber.'
        """
        tables = []
        table_contexts = []
        return tables, table_contexts

    ################################################################
    #                      TABLE HELPERS
    ################################################################
    def _classify_table(self, table: Table) -> str:
        """Classify table type based on content in first row."""
        if not table.rows:
            return 'empty'
        header_text = ' '.join(cell.text.lower() for cell in table.rows[0].cells)
        
        if any(word in header_text for word in ['configuration', 'customizing', 'settings']):
            return 'configuration'
        elif any(word in header_text for word in ['wricef', 'gap', 'requirement']):
            return 'requirements'
        elif any(word in header_text for word in ['version', 'date', 'author']):
            return 'revision'
        elif any(word in header_text for word in ['master', 'data']):
            return 'master_data'
        return 'general'

    def _extract_table_content(self, table: Table) -> str:
        """Extract table content in a simple markdown-like format."""
        content_lines = []
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = self.clean_text(cell.text)
                row_cells.append(cell_text)
            content_lines.append(' | '.join(row_cells))
        return '\n'.join(content_lines)

    def _extract_header_row(self, table: Table) -> List[str]:
        if not table.rows:
            return []
        return [self.clean_text(cell.text) for cell in table.rows[0].cells]

    def _extract_table_context(self, table: Table, table_index: int) -> str:
        """Construct a simple context string for the table."""
        context = f"Table {table_index + 1}"
        if not table.rows:
            return context
        
        header_text = ' '.join(cell.text for cell in table.rows[0].cells)
        if ':' in header_text:
            context += f": {header_text.split(':', 1)[1].strip()}"
        return context

    ################################################################
    #              CHUNKING & KEYWORD EXTRACTION
    ################################################################
    def split_into_chunks(self, text: str, chunk_size: int = 1000) -> List[str]:
        """Split text, using some SAP-specific boundaries if needed."""
        sap_boundaries = [
            r'(?=\b(?:FI|CO|MM|SD|PP|QM|PM|HR|PS)-\d+)',  # Module codes
            r'(?=\bT-Code:)',                             # Transaction codes
            r'(?=\bWRICEF\b)',                            # WRICEF items
            r'(?=\bConfiguration:)',                      # Config sections
            r'(?=\bSolution:)'                            # Etc.
        ]
        
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=200,
            separators=["\n\n", "\n", ". "] + sap_boundaries,
            length_function=len
        )
        
        chunks = splitter.split_text(text)
        return [c for c in chunks if c.strip()]

    def extract_keywords(self, chunks: List[str]) -> List[List[str]]:
        """Use TF-IDF to find top keywords/bigrams for each chunk."""
        if not chunks:
            return []
        try:
            tfidf_matrix = self.vectorizer.fit_transform(chunks)
            feature_names = self.vectorizer.get_feature_names_out()
            
            all_keywords = []
            for row_idx in range(tfidf_matrix.shape[0]):
                row_tfidf = tfidf_matrix[row_idx].toarray()[0]
                top_indices = row_tfidf.argsort()[-10:][::-1]
                keywords = [
                    feature_names[idx]
                    for idx in top_indices
                    if row_tfidf[idx] > 0
                ]
                all_keywords.append(keywords)
            return all_keywords
        except Exception as e:
            self.logger.error(f"Error extracting keywords: {str(e)}")
            return [[] for _ in chunks]

    ################################################################
    #               MAIN PROCESSING PIPELINE
    ################################################################
    def process_document(self, filepath: str) -> Dict:
        """
        Main pipeline:
          - Decide if docx or pdf
          - Extract sections (with page # if docx has manual breaks)
          - Extract tables
          - Chunk each section’s content & assign keywords
          - Return final JSON-friendly structure
        """
        ext = filepath.lower().split('.')[-1]

        if ext == 'docx':
            sections = self.extract_sections_docx(filepath)
            tables, table_contexts = self.extract_tables_docx(filepath)
        elif ext == 'pdf':
            sections = self.extract_sections_pdf(filepath)
            tables, table_contexts = self.extract_tables_pdf(filepath)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        processed_sections = []
        for section in sections:
            # Split content into smaller chunks
            chunks = self.split_into_chunks(section['content'])
            keywords_for_chunks = self.extract_keywords(chunks)

            # Create a separate sub-section for each chunk
            for idx, chunk in enumerate(chunks):
                new_section = section.copy()
                new_section['content'] = chunk
                new_section['keywords'] = (
                    keywords_for_chunks[idx] if idx < len(keywords_for_chunks) else []
                )
                processed_sections.append(new_section)

        processed_doc = {
            'sections': processed_sections,
            'tables': tables,
            'table_contexts': table_contexts,
            'document_name': os.path.basename(filepath)
        }
        return processed_doc

    def _extract_sap_module(self, text: str) -> Optional[str]:
        """Extract SAP module (FI, CO, etc.) from text."""
        module_match = re.search(r'\b(FI|CO|MM|SD|PP|QM|PM|HR|PS)\b', text)
        return module_match.group(1) if module_match else None

################################################################
#        FUNCTION TO PROCESS ALL FILES IN A GIVEN FOLDER
################################################################
def preprocess_all_documents(raw_data_dir="data/raw", preprocessed_dir="data/preprocessed"):
    """Process all .docx/.pdf documents in raw_data_dir, save JSON to preprocessed_dir."""
    os.makedirs(preprocessed_dir, exist_ok=True)
    processor = DocumentPreprocessor()
    
    for filename in os.listdir(raw_data_dir):
        filepath = os.path.join(raw_data_dir, filename)
        if os.path.isdir(filepath):
            continue
        if not (filename.lower().endswith(".docx") or filename.lower().endswith(".pdf")):
            continue
        
        logger.info(f"Processing document: {filename}")
        
        try:
            processed_doc = processor.process_document(filepath)
            
            # Save to JSON
            out_name = os.path.splitext(filename)[0] + ".json"
            output_path = os.path.join(preprocessed_dir, out_name)
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(processed_doc, f, ensure_ascii=False, indent=2)
            
            logger.info(f"Successfully processed and saved: {filename}")
        except Exception as e:
            logger.error(f"Failed to process {filename}: {str(e)}")

if __name__ == "__main__":
    preprocess_all_documents()
